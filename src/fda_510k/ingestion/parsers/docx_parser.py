from pathlib import Path

from docx import Document

from fda_510k.ingestion.parsers.base import BaseParser, ParsedPage, ParserResult


class DocxParser(BaseParser):
    extensions = (".docx",)

    def parse(self, path: Path) -> ParserResult:
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return ParserResult(pages=[ParsedPage(page_number=None, text="\n".join(paragraphs))])
